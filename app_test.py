import streamlit as st
import pandas as pd
import requests
from datetime import datetime
import google.generativeai as genai
import pypdf
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json, io, re, os

# ==============================================================================
# 1. CẤU HÌNH ĐĂNG NHẬP / ĐĂNG KÝ & KẾT NỐI GOOGLE SHEET
# ==============================================================================
SHEET_CSV_URL = "https://docs.google.com/spreadsheets/d/1YHUgWJs3ZNH_6MVYI2Kwowsh7r0XVYaCXopvw1aD0FU/export?format=csv&gid=901150668"
WEB_APP_URL = "https://script.google.com/macros/s/AKfycbzWB6-PRwFkezGzSjS29lrNBVnf03Dy0W1P4S0iDjJ9pIqgD5mDa-qKtc4NTw--IWoPgg/exec"
RULES_FILE = "ai_rules_memory.json"

def load_ai_rules():
    if os.path.exists(RULES_FILE):
        try:
            with open(RULES_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            return []
    return []

def save_ai_rules(rules):
    try:
        with open(RULES_FILE, "w", encoding="utf-8") as f:
            json.dump(rules, f, ensure_ascii=False, indent=2)
        return True
    except:
        return False

def check_login():
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False
    if "user_info" not in st.session_state:
        st.session_state.user_info = {}
    if "reg_missing" not in st.session_state:
        st.session_state.reg_missing = []

    if not st.session_state.logged_in:
        _, center_col, _ = st.columns([1, 1.2, 1])
        with center_col:
            st.markdown("""
            <div style="text-align: center; margin-top: 20px; margin-bottom: 20px;">
                <span style="font-size: 42px;">🔐</span>
                <h2 style="margin: 8px 0 0 0; font-size: 24px; font-weight: 800;">ĐĂNG NHẬP HỆ THỐNG</h2>
                <p style="color: #888; font-size: 13px; margin-top: 4px;">Phần mềm Cụ thể hóa Văn bản Hành chính</p>
            </div>
            """, unsafe_allow_html=True)

            tab_login, tab_register, tab_forgot = st.tabs(["Đăng nhập", "Đăng ký tài khoản", "Quên mật khẩu"])

            # 1. TAB ĐĂNG NHẬP
            with tab_login:
                with st.form("login_form"):
                    username = st.text_input("Tên đăng nhập")
                    password = st.text_input("Mật khẩu", type="password")
                    btn_login = st.form_submit_button("Đăng nhập", use_container_width=True)

                if btn_login:
                    if username == "admin" and password == "Adminai":
                        st.success("Xin chào Admin!")
                        st.session_state.logged_in = True
                        st.session_state.user_info = {"username": "admin", "fullname": "Quản trị viên", "email_phone": "N/A"}
                        st.rerun()
                    else:
                        try:
                            df = pd.read_csv(SHEET_CSV_URL, skiprows=2)
                            df.columns = [c.strip() for c in df.columns]
                            user_match = df[(df['username'].astype(str) == username) & (df['password'].astype(str) == password)]
                            
                            if not user_match.empty:
                                user_data = user_match.iloc[0]
                                st.success(f"Đăng nhập thành công! Chào mừng {user_data['fullname']}")
                                st.session_state.logged_in = True
                                
                                contact_val = "Chưa cập nhật"
                                if 'email_phone' in user_data and pd.notna(user_data['email_phone']):
                                    contact_val = str(user_data['email_phone'])
                                elif 'Email/SĐT' in user_data and pd.notna(user_data['Email/SĐT']):
                                    contact_val = str(user_data['Email/SĐT'])

                                st.session_state.user_info = {
                                    "username": username,
                                    "fullname": user_data['fullname'],
                                    "email_phone": contact_val
                                }
                                st.rerun()
                            else:
                                st.error("Tên đăng nhập hoặc mật khẩu không chính xác!")
                        except Exception as e:
                            st.error("Chưa thể kết nối đến dữ liệu tài khoản.")

            # 2. TAB ĐĂNG KÝ
            with tab_register:
                lbl_user = "Tên đăng nhập mới" + (" :red[*]" if "user" in st.session_state.reg_missing else "")
                lbl_name = "Họ và tên" + (" :red[*]" if "name" in st.session_state.reg_missing else "")
                lbl_contact = "Email hoặc Số điện thoại" + (" :red[*]" if "contact" in st.session_state.reg_missing else "")
                lbl_pass = "Mật khẩu mới" + (" :red[*]" if "pass" in st.session_state.reg_missing else "")
                lbl_conf = "Xác nhận mật khẩu" + (" :red[*]" if "conf" in st.session_state.reg_missing else "")

                with st.form("register_form"):
                    new_user = st.text_input(lbl_user, key="reg_user")
                    new_name = st.text_input(lbl_name, key="reg_name")
                    new_contact = st.text_input(lbl_contact, key="reg_contact")
                    new_pass = st.text_input(lbl_pass, type="password", key="reg_pass")
                    confirm_pass = st.text_input(lbl_conf, type="password", key="reg_conf")
                    btn_register = st.form_submit_button("Đăng ký", use_container_width=True)

                if btn_register:
                    missing = []
                    if not new_user.strip(): missing.append("user")
                    if not new_name.strip(): missing.append("name")
                    if not new_contact.strip(): missing.append("contact")
                    if not new_pass.strip(): missing.append("pass")
                    if not confirm_pass.strip(): missing.append("conf")

                    st.session_state.reg_missing = missing

                    if missing:
                        st.warning("Vui lòng điền đầy đủ các thông tin có dấu (*) đỏ!")
                        st.rerun()
                    elif new_pass != confirm_pass:
                        st.error("Mật khẩu xác nhận không khớp!")
                    else:
                        st.session_state.reg_missing = []
                        payload = {
                            "action": "register",
                            "username": new_user,
                            "password": new_pass,
                            "fullname": new_name,
                            "email_phone": new_contact,
                            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                        try:
                            res = requests.post(WEB_APP_URL, json=payload)
                            if res.status_code == 200:
                                st.success("Đăng ký thành công! Vui lòng quay lại tab Đăng nhập.")
                            else:
                                st.error("Lỗi khi tạo tài khoản.")
                        except:
                            st.error("Không thể kết nối máy chủ đăng ký.")

            # 3. TAB QUÊN MẬT KHẨU
            with tab_forgot:
                with st.form("forgot_form"):
                    fg_user = st.text_input("Nhập Tên đăng nhập của bạn", key="fg_u")
                    fg_contact = st.text_input("Nhập Email hoặc Số điện thoại đã đăng ký", key="fg_c")
                    fg_new_pass = st.text_input("Mật khẩu mới", type="password", key="fg_p1")
                    fg_confirm_pass = st.text_input("Xác nhận mật khẩu mới", type="password", key="fg_p2")
                    btn_forgot = st.form_submit_button("Đặt lại mật khẩu", use_container_width=True)

                if btn_forgot:
                    if not fg_user or not fg_contact or not fg_new_pass:
                        st.warning("Vui lòng điền đầy đủ thông tin!")
                    elif fg_new_pass != fg_confirm_pass:
                        st.error("Mật khẩu xác nhận không khớp!")
                    else:
                        try:
                            df = pd.read_csv(SHEET_CSV_URL, skiprows=2)
                            df.columns = [c.strip() for c in df.columns]
                            col_email = 'email_phone' if 'email_phone' in df.columns else 'Email/SĐT'
                            match = df[(df['username'].astype(str) == fg_user) & (df[col_email].astype(str) == fg_contact)]
                            
                            if not match.empty:
                                payload = {
                                    "action": "update_password",
                                    "username": fg_user,
                                    "new_password": fg_new_pass
                                }
                                res = requests.post(WEB_APP_URL, json=payload)
                                if res.status_code == 200:
                                    st.success("Đổi mật khẩu thành công! Vui lòng quay lại tab Đăng nhập.")
                                else:
                                    st.error("Lỗi khi cập nhật mật khẩu.")
                            else:
                                st.error("Tên đăng nhập và Email/SĐT không khớp!")
                        except Exception as e:
                            st.error("Không thể xác minh thông tin.")

        return False
    return True

# Thiết lập Cấu hình Trang Streamlit
st.set_page_config(page_title="Phần mềm Cụ thể hóa Văn bản Hành chính", page_icon="🏛️", layout="wide")

# Bắt buộc Đăng nhập
if not check_login():
    st.stop()

# ==============================================================================
# 2. CẤU HÌNH GIAO DIỆN & BẢNG ÁNH XẠ
# ==============================================================================
SYMBOL_MAP = {
    "Kế hoạch": "KH",
    "Công văn": "CV",
    "Báo cáo": "BC",
    "Tờ trình": "TTr",
    "Thông báo": "TB",
    "Quyết định": "QĐ"
}

a4_css = """
<style>
.app-header {
    background: linear-gradient(135deg, #7b0000 0%, #a81010 50%, #c41e1e 100%);
    border: 1px solid #e0a800;
    border-radius: 12px;
    padding: 20px 25px;
    margin-bottom: 20px;
    box-shadow: 0 4px 15px rgba(0, 0, 0, 0.4);
    display: flex;
    align-items: center;
    justify-content: space-between;
}
.app-header-title {
    color: #ffffff;
    font-size: 24px;
    font-weight: bold;
    letter-spacing: 0.5px;
    text-shadow: 1px 1px 3px rgba(0,0,0,0.6);
    margin: 0;
}
.app-header-sub {
    color: #ffd700;
    font-size: 13px;
    margin-top: 4px;
    font-weight: 500;
}
.section-badge {
    background: linear-gradient(90deg, #d4af37 0%, #f3e5ab 100%);
    color: #4a2c00;
    font-size: 11px;
    font-weight: bold;
    padding: 3px 8px;
    border-radius: 4px;
    text-transform: uppercase;
    display: inline-block;
    margin-bottom: 6px;
}
.a4-wrapper {
    background: radial-gradient(circle, #3d434d 0%, #20242b 100%);
    padding: 20px;
    border-radius: 10px;
    border: 1px solid #4a5568;
    display: flex;
    justify-content: center;
    width: 100%;
    box-shadow: inset 0 0 15px rgba(0,0,0,0.5);
}
.a4-paper {
    background-color: #ffffff !important;
    color: #000000 !important;
    width: 100%;
    padding: 30px 35px;
    font-family: 'Times New Roman', Times, serif;
    font-size: 10.5pt;
    line-height: 1.35;
    box-shadow: 0 10px 25px rgba(0, 0, 0, 0.6);
    max-height: 620px;
    overflow-y: auto;
    box-sizing: border-box;
    border-radius: 2px;
}
.header-table, .footer-table {
    width: 100%;
    border-collapse: collapse;
    margin-bottom: 10px;
    border: none !important;
}
.header-table td, .footer-table td {
    vertical-align: top;
    font-family: 'Times New Roman', Times, serif;
    font-size: 10pt;
    line-height: 1.2;
    color: #000000;
    padding: 0px;
}
.custom-underline {
    display: inline-block;
    border-bottom: 1px solid #000000;
    padding-bottom: 2px;
    line-height: 1.1;
}
.title-block {
    text-align: center;
    font-weight: bold;
    font-size: 12pt;
    margin-top: 10px;
    margin-bottom: 4px;
}
.trich-yeu-block {
    text-align: center;
    font-weight: bold;
    font-size: 11pt;
    margin-top: 4px;
    margin-bottom: 4px;
}
.short-line {
    width: 35%;
    margin: 6px auto 12px auto;
    border: 0;
    border-top: 1px solid #000000;
}
.content-para {
    text-align: justify;
    text-indent: 1cm;
    margin-bottom: 5px;
    line-height: 1.35;
}
.heading-para {
    font-weight: bold;
    font-size: 10.5pt;
    margin-top: 10px;
    margin-bottom: 3px;
}
.noi-nhan-block {
    padding-left: 1cm !important;
    text-align: left;
    font-size: 9.5pt;
    line-height: 1.2;
}
.chat-user-box {
    background: linear-gradient(90deg, #242933 0%, #1a1e24 100%);
    color: #ffffff;
    padding: 12px 16px;
    border-radius: 8px;
    margin-bottom: 8px;
    display: flex;
    align-items: center;
    font-size: 13px;
    border-left: 3px solid #e53e3e;
    border: 1px solid #323946;
}
.chat-user-icon {
    background: linear-gradient(135deg, #e53e3e 0%, #9b2c2c 100%);
    color: white;
    width: 30px;
    height: 30px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    margin-right: 12px;
    font-size: 14px;
    flex-shrink: 0;
    box-shadow: 0 2px 5px rgba(0,0,0,0.3);
}
.chat-ai-box {
    background: linear-gradient(90deg, #1f2d24 0%, #17241c 100%);
    color: #ffffff;
    padding: 12px 16px;
    border-radius: 8px;
    margin-bottom: 16px;
    display: flex;
    align-items: center;
    font-size: 13px;
    border-left: 3px solid #38a169;
    border: 1px solid #234e32;
}
.chat-ai-icon {
    background: linear-gradient(135deg, #38a169 0%, #22543d 100%);
    color: white;
    width: 30px;
    height: 30px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    margin-right: 12px;
    font-size: 14px;
    flex-shrink: 0;
    box-shadow: 0 2px 5px rgba(0,0,0,0.3);
}
</style>
"""
st.markdown(a4_css, unsafe_allow_html=True)

# --- THÔNG TIN TÀI KHOẢN TRÊN SIDEBAR ---
with st.sidebar:
    st.markdown("""
    <div style="text-align: center; padding-bottom: 10px;">
        <span style="font-size: 32px;">🏛️</span>
        <h3 style="color: #ffd700; margin: 0; font-size: 18px;">HỆ THỐNG ĐIỀU HÀNH</h3>
        <p style="color: #a0aec0; font-size: 11px;">Chuẩn Thể Thức Đảng & Nhà Nước</p>
    </div>
    """, unsafe_allow_html=True)
    st.write("---")
    
    the_thuc = st.radio("Chọn Khối văn bản:", ["Khối Đảng", "Khối Nhà nước"])
    if the_thuc == "Khối Đảng":
        st.info("📌 **Áp dụng:** Hướng dẫn 05-HD/VPTW của Văn phòng Trung ương Đảng")
    else:
        st.info("📌 **Áp dụng:** Nghị định 30/2020/NĐ-CP của Chính phủ")
        
    st.subheader("⚙️ Cấu hình AI")
    ai_provider = st.selectbox("AI xử lý chính", ["Google Gemini"])
    api_key = st.text_input("Gemini API key", type="password")
    model_name = st.selectbox("Model", ["gemini-3.6-flash"])

    # KHU VỰC QUẢN LÝ QUY TẮC GHI NHỚ
    st.write("---")
    st.subheader("🧠 Sổ tay ghi nhớ AI")
    current_rules = load_ai_rules()
    if current_rules:
        st.caption("Các quy tắc AI đang ghi nhớ và áp dụng:")
        for idx, r in enumerate(current_rules):
            st.markdown(f"• {r}")
        if st.button("🗑️ Xóa toàn bộ ghi nhớ", use_container_width=True):
            save_ai_rules([])
            st.success("Đã xóa bộ nhớ quy tắc!")
            st.rerun()
    else:
        st.caption("Chưa có quy tắc ghi nhớ nào. Khi chat sửa đổi, anh có thể dặn 'Nhớ luôn quy tắc...' để AI tự lưu.")

    st.write("---")
    user_info = st.session_state.get("user_info", {})
    with st.popover(f"👤 Tài khoản ({user_info.get('fullname', 'User')})"):
        st.markdown(f"**Họ tên:** {user_info.get('fullname')}")
        st.markdown(f"**Username:** {user_info.get('username')}")
        st.markdown(f"**Liên hệ:** {user_info.get('email_phone')}")
        st.write("---")
        if st.button("🚪 Đăng xuất", use_container_width=True):
            st.session_state.logged_in = False
            st.session_state.user_info = {}
            st.rerun()

# ==============================================================================
# 3. GIAO DIỆN CHÍNH
# ==============================================================================
st.markdown("""
<div class="app-header">
    <div>
        <div class="app-header-title">🏛️ PHẦN MỀM CỤ THỂ HÓA VĂN BẢN HÀNH CHÍNH</div>
        <div class="app-header-sub">HỆ THỐNG HỖ TRỢ BIÊN SOẠN & XỬ LÝ VĂN KIỆN ĐẢNG - CHÍNH QUYỀN TỰ ĐỘNG BẰNG AI</div>
    </div>
    <div style="font-size: 38px;">🇻🇳</div>
</div>
""", unsafe_allow_html=True)

# MỤC 1 & 2
col1, col2 = st.columns([1, 1])
with col1:
    st.markdown('<span class="section-badge">BƯỚC 1</span> <b>File nguồn & Loại văn bản</b>', unsafe_allow_html=True)
    uploaded_files = st.file_uploader(
        "Tải file nguồn/Đề cương (.docx, .pdf, .png, .jpg...):",
        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
        accept_multiple_files=True
    )
    loai_vb = st.selectbox("Chọn Loại văn bản đầu ra:", ["Kế hoạch", "Công văn", "Báo cáo", "Tờ trình", "Thông báo", "Quyết định"])

with col2:
    st.markdown('<span class="section-badge">BƯỚC 2</span> <b>Yêu cầu & Cơ quan ban hành</b>', unsafe_allow_html=True)
    yeu_cau = st.text_area("Anh muốn cụ thể hóa như thế nào?:", height=100, placeholder="Soạn thảo văn bản theo yêu cầu chỉ đạo...")
    co_quan_default = "ĐẢNG ỦY PHƯƠNG NHƠN TRẠCH" if the_thuc == "Khối Đảng" else "UBND PHƯƠNG NHƠN TRẠCH"
    
    if "current_agency" not in st.session_state:
        st.session_state.current_agency = co_quan_default
    co_quan = st.text_input("Cơ quan ban hành dự thảo:", value=st.session_state.current_agency)

# MỤC 3
st.markdown('<br>', unsafe_allow_html=True)
st.markdown('<span class="section-badge">BƯỚC 3</span> <b>File mẫu riêng & Mẫu gợi ý chuẩn</b>', unsafe_allow_html=True)
col3_1, col3_2 = st.columns([1, 1])
with col3_1:
    custom_template_file = st.file_uploader("Tải file mẫu riêng (Chỉ lấy thể thức/khung mẫu):", type=["docx"], key="custom_template")
with col3_2:
    de_cuong_goy_y = st.selectbox("📌 Mẫu gợi ý / Đề cương chuẩn:", [
        "(Không chọn mẫu gợi ý)",
        "Đề cương chuẩn Hướng dẫn 05-HD/VPTW (Công tác Đảng)",
        "Đề cương Kế hoạch hành động 100 ngày Chuyển đổi số",
        "Đề cương Báo cáo kết quả thực hiện nhiệm vụ chính trị"
    ])

st.markdown("---")
btn_process = st.button("⚡ PHÂN TÍCH & CỤ THỂ HÓA VĂN BẢN", type="primary", use_container_width=True)

if "draft_text" not in st.session_state:
    st.session_state.draft_text = ""
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# --- XỬ LÝ AI TẠO DỰ THẢO ---
if btn_process:
    if not api_key:
        st.error("Vui lòng nhập Gemini API Key ở cột bên trái!")
    elif not uploaded_files and not yeu_cau:
        st.warning("Vui lòng tải file nguồn hoặc nhập yêu cầu!")
    else:
        with st.spinner("Đang phân tích dữ liệu và tổng hợp văn bản chuẩn thể thức..."):
            try:
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(model_name)
                
                content_parts = []
                extracted_texts = []
                
                for uf in uploaded_files:
                    bytes_data = uf.read()
                    if uf.name.lower().endswith('.pdf'):
                        try:
                            reader = pypdf.PdfReader(io.BytesIO(bytes_data))
                            pdf_text = "".join([page.extract_text() or "" for page in reader.pages])
                            if len(pdf_text.strip()) > 50:
                                extracted_texts.append(f"--- NỘI DUNG FILE NGUỒN {uf.name} ---\n" + pdf_text)
                            else:
                                content_parts.append({"mime_type": "application/pdf", "data": bytes_data})
                        except:
                            content_parts.append({"mime_type": "application/pdf", "data": bytes_data})
                    elif uf.name.lower().endswith('.docx'):
                        doc_file = docx.Document(io.BytesIO(bytes_data))
                        docx_text = "\n".join([p.text for p in doc_file.paragraphs])
                        extracted_texts.append(f"--- NỘI DUNG FILE NGUỒN {uf.name} ---\n" + docx_text)
                    else:
                        content_parts.append({"mime_type": uf.type, "data": bytes_data})
                
                custom_template_prompt = ""
                if custom_template_file is not None:
                    try:
                        doc_tpl = docx.Document(io.BytesIO(custom_template_file.read()))
                        tpl_text = "\n".join([p.text for p in doc_tpl.paragraphs if p.text.strip()])
                        if tpl_text:
                            custom_template_prompt = f"\nBÁM SÁT KHUNG MẪU VĂN BẢN ĐÍNH KÈM (Chỉ học theo thể thức, bố cục mục I, II, III và phong cách trình bày):\n--- MẪU THỂ THỨC KHUNG ---\n{tpl_text}\n--- KẾT THÚC MẪU ---\n"
                    except Exception as tpl_err:
                        st.warning(f"Lỗi đọc file mẫu: {str(tpl_err)}")

                de_cuong_prompt = ""
                if de_cuong_goy_y != "(Không chọn mẫu gợi ý)":
                    de_cuong_prompt = f"\nÁP DỤNG ĐỀ CƯƠNG: {de_cuong_goy_y}"

                # NẠP CÁC QUY TẮC ĐÃ GHI NHỚ TỪ TRƯỚC
                saved_rules_prompt = ""
                loaded_rules = load_ai_rules()
                if loaded_rules:
                    rules_str = "\n".join([f"- {r}" for r in loaded_rules])
                    saved_rules_prompt = f"\nCÁC QUY TẮC CỐ ĐỊNH PHẢI TUÂN THỦ TỪ TRƯỚC ĐẾN NAY:\n{rules_str}\n"

                rule_doc_type = ""
                if loai_vb == "Công văn":
                    rule_doc_type = """
                    ĐÂY LÀ CÔNG VĂN HÀNH CHÍNH:
                    1. TUYỆT ĐỐI KHÔNG VIẾT tiêu đề "CÔNG VĂN" căn giữa trang.
                    2. Dòng đầu tiên phải ghi rõ Trích yếu nội dung dạng: "V/v [Tóm tắt nội dung công văn]"
                    3. Ngay sau trích yếu là dòng "Kính gửi: [Tên các cơ quan/đơn vị nhận công văn]".
                    """
                else:
                    rule_doc_type = f"""
                    ĐÂY LÀ VĂN BẢN CÓ TÊN LOẠI ({loai_vb.upper()}):
                    1. Dòng đầu tiên ghi TÊN LOẠI VĂN BẢN viết hoa căn giữa (Ví dụ: {loai_vb.upper()}).
                    2. Dòng tiếp theo ghi TRÍCH YẾU NỘI DUNG của {loai_vb}.
                    """

                prompt = f"""
                Bạn là chuyên gia soạn thảo văn bản hành chính Việt Nam.
                Hãy soạn thảo nội dung của 01 dự thảo văn bản hoàn chỉnh dựa trên tài liệu đính kèm.
                
                THỂ THỨC: {the_thuc} | CƠ QUAN BAN HÀNH: {co_quan} | LOẠI VĂN BẢN: {loai_vb}
                YÊU CẦU CỤ THỂ HÓA: {yeu_cau}
                {de_cuong_prompt}
                {custom_template_prompt}
                {saved_rules_prompt}
                
                DỮ LIỆU TÀI LIỆU NGUỒN:
                {"".join(extracted_texts)}
                
                QUY TẮC THỂ THỨC BẮT BUỘC:
                {rule_doc_type}
                - TUYỆT ĐỐI KHÔNG VIẾT Quốc hiệu, Tiêu ngữ, Tên cơ quan ban hành, Số/Ký hiệu, Ngày tháng ở đầu bài (Vì giao diện đã tự chèn).
                - TUYỆT ĐỐI KHÔNG DÙNG KÝ TỰ MARKDOWN (*, #, _).
                - Nội dung bao gồm Căn cứ pháp lý, Các mục nội dung.
                - Ở MỤC CUỐI CÙNG (NƠI NHẬN & CHỮ KÝ): Bắt buộc ghi dạng:
                  Nơi nhận:
                  - Như trên;
                  - Lưu: VP.
                  T/M BAN THƯỜNG VỤ
                  BÍ THƯ
                  (Để trống khoảng ký tên)
                  Họ và Tên
                """
                content_parts.insert(0, prompt)
                
                response = model.generate_content(content_parts)
                st.session_state.draft_text = response.text
                st.session_state.current_agency = co_quan
                st.session_state.chat_history = []
                st.success("Đã cụ thể hóa văn bản thành công!")
            except Exception as e:
                st.error(f"Lỗi xử lý API: {str(e)}")

# ==============================================================================
# 4. GIAO DIỆN HIỂN THỊ DỰ THẢO A4 & CHAT AI SỬA ĐỔI
# ==============================================================================
if st.session_state.draft_text:
    res_col1, res_col2 = st.columns([1.2, 0.8])
    
    with res_col1:
        st.markdown('##### 📄 BẢN DỰ THẢO VĂN BẢN (A4)')
        
        clean_text = re.sub(r'[\*#_]', '', st.session_state.draft_text)
        raw_lines = [l.strip() for l in clean_text.split('\n') if l.strip()]
        
        filtered_lines = []
        for l in raw_lines:
            l_up = l.upper()
            if any(k in l_up for k in ["ĐẢNG BỘ", "ĐẢNG CỘNG SẢN", "CỘNG HÒA XÃ HỘI", "ĐỘC LẬP - TỰ DO"]) and len(l) < 80:
                continue
            if re.search(r',\s*NGÀY\s+.*\s+THÁNG\s+.*\s+NĂM', l_up) and len(l) < 80:
                continue
            if ("UBND" in l_up or "ĐẢNG ỦY" in l_up) and len(l) < 60 and not l_up.startswith("KẾ HOẠCH") and not l_up.startswith("CÔNG VĂN"):
                continue
            if re.match(r'^SỐ\s*:', l_up) or re.match(r'^SỐ\s*-', l_up):
                continue
            filtered_lines.append(l)

        body_lines = []
        noi_nhan_list = ["- Như trên;", "- Lưu: VP."]
        chuc_vu_signer = "T/M BAN THƯỜNG VỤ\nBÍ THƯ" if the_thuc == "Khối Đảng" else "TM. ỦY BAN NHÂN DÂN\nCHỦ TỊCH"
        ten_signer = "Họ và Tên"
        
        in_footer = False
        
        for l in filtered_lines:
            l_strip = l.strip()
            l_up = l_strip.upper()
            
            if l_up.startswith("NƠI NHẬN:") or l_up == "NƠI NHẬN":
                in_footer = True
                continue
            
            if in_footer:
                if l_strip.startswith("-"):
                    if l_strip not in noi_nhan_list:
                        noi_nhan_list.append(l_strip)
                elif any(k in l_up for k in ["T/M", "TM.", "BÍ THƯ", "CHỦ TỊCH", "PHÓ BÍ THƯ"]):
                    chuc_vu_signer = l_strip
                elif len(l_strip) < 40 and not l_strip.startswith("I") and not l_strip.startswith("1") and l_strip.lower() != "họ và tên":
                    ten_signer = l_strip
            else:
                if any(k in l_up for k in ["T/M ", "TM. ", "BÍ THƯ", "CHỦ TỊCH"]) and len(l_strip) < 50:
                    in_footer = True
                    chuc_vu_signer = l_strip
                else:
                    body_lines.append(l_strip)

        trich_yeu_cv = ""
        if loai_vb == "Công văn" and body_lines:
            if body_lines[0].startswith("V/v") or body_lines[0].startswith("Về việc"):
                trich_yeu_cv = body_lines.pop(0)

        # Tự động đồng bộ địa danh ngày tháng từ tên cơ quan
        agency_display = st.session_state.get("current_agency", co_quan)
        dia_danh = "Nhơn Trạch"
        if "ĐẠI PHƯỚC" in agency_display.upper():
            dia_danh = "Đại Phước"
        elif "NHƠN TRẠCH" in agency_display.upper():
            dia_danh = "Nhơn Trạch"
        else:
            match_dd = re.search(r'(phường|xã|thị trấn|huyện|thành phố)\s+([^\n,]+)', agency_display, re.IGNORECASE)
            if match_dd:
                dia_danh = match_dd.group(2).strip().title()

        type_code = SYMBOL_MAP.get(loai_vb, "CV")
        so_ky_hieu = f"-{type_code}/ĐU" if the_thuc == "Khối Đảng" else f"/{type_code}-UBND"

        if the_thuc == "Khối Đảng":
            sub_cv = f"<br><br><i>{trich_yeu_cv}</i>" if trich_yeu_cv else ""
            header_table = f'<table class="header-table"><tr><td style="width: 48%; text-align: center;"><b>ĐẢNG BỘ THÀNH PHỐ ĐỒNG NAI</b><br><b>{agency_display.upper()}</b><br><span style="font-size: 7pt;">*</span><br>Số: &nbsp;&nbsp;&nbsp;&nbsp;{so_ky_hieu}{sub_cv}</td><td style="width: 52%; text-align: center;"><span class="custom-underline"><b>ĐẢNG CỘNG SẢN VIỆT NAM</b></span><br><br><i>{dia_danh}, ngày &nbsp;&nbsp;&nbsp; tháng 8 năm 2026</i></td></tr></table>'
        else:
            sub_cv = f"<br><br><i>{trich_yeu_cv}</i>" if trich_yeu_cv else ""
            header_table = f'<table class="header-table"><tr><td style="width: 45%; text-align: center;">UBND THÀNH PHỐ ĐỒNG NAI<br><b><span class="custom-underline">{agency_display.upper()}</span></b><br><span style="font-size: 7pt;">*</span><br>Số: &nbsp;&nbsp;&nbsp;&nbsp;{so_ky_hieu}{sub_cv}</td><td style="width: 55%; text-align: center;"><b>CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM</b><br><b><span class="custom-underline">Độc lập - Tự do - Hạnh phúc</span></b><br><i>{dia_danh}, ngày &nbsp;&nbsp;&nbsp; tháng 8 năm 2026</i></td></tr></table>'

        body_content = ""
        is_trich_yeu = False
        
        for line in body_lines:
            if re.match(r'^(I|II|III|IV|V|VI|VII|VIII)\.', line):
                body_content += f'<div class="heading-para">{line}</div>'
                is_trich_yeu = False
            elif re.match(r'^\d+\.', line) and len(line) < 80:
                body_content += f'<div class="heading-para">{line}</div>'
                is_trich_yeu = False
            elif line.isupper() and len(line) < 100 and loai_vb != "Công văn":
                body_content += f'<div class="title-block">{line}</div>'
                is_trich_yeu = True
            elif is_trich_yeu and loai_vb != "Công văn":
                body_content += f'<div class="trich-yeu-block">{line}</div><hr class="short-line">'
                is_trich_yeu = False
            elif line.startswith("Kính gửi:") or line.startswith("-"):
                body_content += f'<div style="text-align: left; margin-bottom: 4px; padding-left: 10px;">{line}</div>'
                is_trich_yeu = False
            else:
                body_content += f'<div class="content-para">{line}</div>'
                is_trich_yeu = False

        noi_nhan_html = "<br>".join(noi_nhan_list)
        chuc_vu_html = chuc_vu_signer.replace("\n", "<br>")
        footer_table = f"""
        <table class="footer-table" style="margin-top: 20px;">
            <tr>
                <td style="width: 45%; text-align: left; vertical-align: top;" class="noi-nhan-block">
                    <b><u>Nơi nhận:</u></b><br>
                    {noi_nhan_html}
                </td>
                <td style="width: 55%; text-align: center; vertical-align: top;">
                    <b>{chuc_vu_html}</b><br><br><br><br><br>
                    <b>{ten_signer}</b>
                </td>
            </tr>
        </table>
        """

        full_a4_html = f'<div class="a4-wrapper"><div class="a4-paper">{header_table}{body_content}{footer_table}</div></div>'
        st.markdown(full_a4_html, unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

        def generate_docx(b_lines, agency_name, form_type, doc_type_str, cv_subj, n_nhan, c_vu, t_ky, dia_danh_str):
            doc = docx.Document()
            for section in doc.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(3)
                section.right_margin = Cm(2)
            
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            
            cell_left, cell_right = table.cell(0, 0), table.cell(0, 1)
            cell_left.width, cell_right.width = Cm(8.5), Cm(8.5)
            
            t_code = SYMBOL_MAP.get(doc_type_str, "CV")
            code_str = f"-{t_code}/ĐU" if form_type == "Khối Đảng" else f"/{t_code}-UBND"

            p_left = cell_left.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_left.paragraph_format.line_spacing = 1.15
            p_left.paragraph_format.space_after = Pt(0)
            
            if form_type == "Khối Đảng":
                r1 = p_left.add_run("ĐẢNG BỘ THÀNH PHỐ ĐỒNG NAI\n")
                r1.font.name, r1.font.size = 'Times New Roman', Pt(12)
                r2 = p_left.add_run(f"{agency_name.upper()}\n")
                r2.font.name, r2.font.size, r2.font.bold = 'Times New Roman', Pt(12), True
                r3 = p_left.add_run("*\n")
                r3.font.name, r3.font.size = 'Times New Roman', Pt(9)
                r4 = p_left.add_run(f"Số:       {code_str}")
                r4.font.name, r4.font.size = 'Times New Roman', Pt(12)
                if cv_subj:
                    r5 = p_left.add_run(f"\n\n{cv_subj}")
                    r5.font.name, r5.font.size, r5.font.italic = 'Times New Roman', Pt(11), True
            else:
                r1 = p_left.add_run("UBND THÀNH PHỐ ĐỒNG NAI\n")
                r1.font.name, r1.font.size = 'Times New Roman', Pt(12)
                r2 = p_left.add_run(f"{agency_name.upper()}\n")
                r2.font.name, r2.font.size, r2.font.bold, r2.font.underline = 'Times New Roman', Pt(12), True, True
                r3 = p_left.add_run("*\n")
                r3.font.name, r3.font.size = 'Times New Roman', Pt(9)
                r4 = p_left.add_run(f"Số:       {code_str}")
                r4.font.name, r4.font.size = 'Times New Roman', Pt(12)
                if cv_subj:
                    r5 = p_left.add_run(f"\n\n{cv_subj}")
                    r5.font.name, r5.font.size, r5.font.italic = 'Times New Roman', Pt(11), True

            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_right.paragraph_format.line_spacing = 1.15
            p_right.paragraph_format.space_after = Pt(0)

            if form_type == "Khối Đảng":
                r1 = p_right.add_run("ĐẢNG CỘNG SẢN VIỆT NAM")
                r1.font.name, r1.font.size, r1.font.bold, r1.font.underline = 'Times New Roman', Pt(12), True, True
                r2 = p_right.add_run(f"\n\n{dia_danh_str}, ngày     tháng 8 năm 2026")
                r2.font.name, r2.font.size, r2.font.italic = 'Times New Roman', Pt(12), True
            else:
                r1 = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
                r1.font.name, r1.font.size, r1.font.bold = 'Times New Roman', Pt(12), True
                r2 = p_right.add_run("Độc lập - Tự do - Hạnh phúc")
                r2.font.name, r2.font.size, r2.font.bold, r2.font.underline = 'Times New Roman', Pt(12.5), True, True
                r3 = p_right.add_run(f"\n{dia_danh_str}, ngày     tháng 8 năm 2026")
                r3.font.name, r3.font.size, r3.font.italic = 'Times New Roman', Pt(12), True

            doc.add_paragraph().paragraph_format.space_after = Pt(6)

            next_is_trich_yeu = False
            for line in b_lines:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.2
                p.paragraph_format.space_after = Pt(4)
                
                if re.match(r'^(I|II|III|IV|V|VI|VII|VIII)\.', line):
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(line)
                    run.font.name, run.font.size, run.font.bold = 'Times New Roman', Pt(13), True
                    next_is_trich_yeu = False
                elif line.isupper() and doc_type_str != "Công văn":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line)
                    run.font.name, run.font.size, run.font.bold = 'Times New Roman', Pt(14), True
                    next_is_trich_yeu = True
                elif next_is_trich_yeu and doc_type_str != "Công văn":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line)
                    run.font.name, run.font.size, run.font.bold = 'Times New Roman', Pt(13), True
                    
                    p_line = doc.add_paragraph()
                    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_line.paragraph_format.space_after = Pt(8)
                    r_line = p_line.add_run("                  ")
                    r_line.font.name, r_line.font.size, r_line.font.underline = 'Times New Roman', Pt(12), True
                    next_is_trich_yeu = False
                elif re.match(r'^\d+\.', line) and len(line) < 80:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(line)
                    run.font.name, run.font.size, run.font.bold = 'Times New Roman', Pt(13), True
                    next_is_trich_yeu = False
                elif line.startswith("Kính gửi:"):
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(line)
                    run.font.name, run.font.size = 'Times New Roman', Pt(13)
                    next_is_trich_yeu = False
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(1.27)
                    run = p.add_run(line)
                    run.font.name, run.font.size = 'Times New Roman', Pt(13)
                    next_is_trich_yeu = False

            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            t_foot = doc.add_table(rows=1, cols=2)
            t_foot.alignment = WD_TABLE_ALIGNMENT.CENTER
            t_foot.autofit = False
            
            c_f_left, c_f_right = t_foot.cell(0, 0), t_foot.cell(0, 1)
            c_f_left.width, c_f_right.width = Cm(8.5), Cm(8.5)

            p_f_l = c_f_left.paragraphs[0]
            p_f_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_f_l.paragraph_format.line_spacing = 1.15
            p_f_l.paragraph_format.left_indent = Cm(1)
            
            r_nn_title = p_f_l.add_run("Nơi nhận:\n")
            r_nn_title.font.name, r_nn_title.font.size, r_nn_title.font.bold, r_nn_title.font.underline, r_nn_title.font.italic = 'Times New Roman', Pt(11), True, True, True
            
            for item in n_nhan:
                r_item = p_f_l.add_run(f"{item}\n")
                r_item.font.name, r_item.font.size = 'Times New Roman', Pt(11)

            p_f_r = c_f_right.paragraphs[0]
            p_f_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_f_r.paragraph_format.line_spacing = 1.15
            
            r_cv = p_f_r.add_run(f"{c_vu}\n\n\n\n\n")
            r_cv.font.name, r_cv.font.size, r_cv.font.bold = 'Times New Roman', Pt(13), True
            
            r_ten = p_f_r.add_run(t_ky)
            r_ten.font.name, r_ten.font.size, r_ten.font.bold = 'Times New Roman', Pt(13), True

            bio = io.BytesIO()
            doc.save(bio)
            return bio.getvalue()

        st.download_button(
            label="📥 TẢI VỀ FILE WORD (.DOCX)",
            data=generate_docx(body_lines, agency_display, the_thuc, loai_vb, trich_yeu_cv, noi_nhan_list, chuc_vu_signer, ten_signer, dia_danh),
            file_name=f"Du_Thao_{loai_vb}_{dia_danh}.docx".replace(" ", "_"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )

    # --- KHU VỰC CHAT AI SỬA ĐỔI ---
    with res_col2:
        st.markdown('##### 💬 TRỢ LÝ AI CHỈNH SỬA (GOOGLE GEMINI)')
        st.caption("Nhập yêu cầu (VD: 'Sửa căn cứ 1', 'Bỏ mục II') để AI cập nhật trực tiếp lên trang Word bên trái.")
        
        edit_instruction = st.text_area("Nhập yêu cầu chỉnh sửa...", height=120, label_visibility="collapsed", placeholder="Nhập yêu cầu chỉnh sửa văn bản...")
        
        if st.button("Chỉnh sửa dự thảo", use_container_width=True):
            if edit_instruction and api_key:
                with st.spinner("AI đang cập nhật lại dự thảo..."):
                    try:
                        genai.configure(api_key=api_key)
                        model = genai.GenerativeModel(model_name)
                        
                        # Nạp các quy tắc đã có để AI tham chiếu khi sửa
                        saved_rules = load_ai_rules()
                        rules_ctx = "\n".join([f"- {r}" for r in saved_rules]) if saved_rules else "Chưa có quy tắc riêng."
                        
                        edit_prompt = f"""
                        BẢN DỰ THẢO HIỆN TẠI:
                        {st.session_state.draft_text}

                        CÁC QUY TẮC ĐÃ LƯU TRƯỚC ĐÂY:
                        {rules_ctx}

                        YÊU CẦU CHỈNH SỬA TỪ NGƯỜI DÙNG:
                        {edit_instruction}

                        HÃY CẬP NHẬT LẠI TOÀN BỘ BẢN DỰ THẢO THEO ĐÚNG YÊU CẦU TRÊN.
                        TUYỆT ĐỐI KHÔNG DÙNG KÝ TỰ MARKDOWN (*, #, _).
                        """
                        res_edit = model.generate_content(edit_prompt)
                        st.session_state.draft_text = res_edit.text
                        
                        # Tự động cập nhật tên cơ quan nếu người dùng yêu cầu đổi địa danh/cơ quan
                        edit_lower = edit_instruction.lower()
                        if "đại phước" in edit_lower:
                            st.session_state.current_agency = st.session_state.current_agency.replace("NHƠN TRẠCH", "ĐẠI PHƯỚC").replace("Nhơn Trạch", "Đại Phước")
                        elif "nhơn trạch" in edit_lower:
                            st.session_state.current_agency = st.session_state.current_agency.replace("ĐẠI PHƯỚC", "NHƠN TRẠCH").replace("Đại Phước", "Nhơn Trạch")
                        
                        # KIỂM TRA VÀ TỰ ĐỘNG LƯU VÀO BỘ NHỚ QUY TẮC
                        keywords_remember = ["nhớ", "lưu ý", "từ nay", "các văn bản sau", "luôn luôn", "quy tắc"]
                        if any(k in edit_lower for k in keywords_remember):
                            if edit_instruction not in saved_rules:
                                saved_rules.append(edit_instruction)
                                save_ai_rules(saved_rules)
                        
                        st.session_state.chat_history.append(edit_instruction)
                        st.rerun()
                    except Exception as e:
                        err_str = str(e)
                        if "429" in err_str or "quota" in err_str.lower():
                            st.warning("⚠️ API Key của anh đang tạm hết lượt gọi trong ngày (Rate Limit). Anh vui lòng đợi khoảng 1-2 phút rồi bấm lại, hoặc đổi Key Gemini khác ở cột bên trái nhé!")
                        else:
                            st.error(f"Lỗi: {err_str}")

        if st.session_state.chat_history:
            st.markdown("<br>", unsafe_allow_html=True)
            for cmd in reversed(st.session_state.chat_history):
                user_html = f'<div class="chat-user-box"><div class="chat-user-icon">⏰</div><div>{cmd}</div></div>'
                ai_html = '<div class="chat-ai-box"><div class="chat-ai-icon">📊</div><div><b>✅ Đã cập nhật văn bản lên trang Word!</b></div></div>'
                st.markdown(user_html + ai_html, unsafe_allow_html=True)
