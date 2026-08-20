import streamlit as st
import pandas as pd
import requests
from datetime import datetime
import google.generativeai as genai
import pypdf
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json, io, re, os

SHEET_CSV_URL = "https://docs.google.com/spreadsheets/d/1YHUgWJs3ZNH_6MVYI2Kwowsh7r0XVYaCXopvw1aD0FU/export?format=csv&gid=901150668"
WEB_APP_URL = "https://script.google.com/macros/s/AKfycbzWB6-PRwFkezGzSjS29lrNBVnf03Dy0W1P4S0iDjJ9pIqgD5mDa-qKtc4NTw--IWoPgg/exec"
CONFIG_FILE = "config_keys.json"

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f: return json.load(f)
        except Exception: return {}
    return {}

def save_config(data):
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f: json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception: return False

config_data = load_config()

def check_login():
    if "logged_in" not in st.session_state: st.session_state.logged_in = False
    if "user_info" not in st.session_state: st.session_state.user_info = {}
    if "reg_missing" not in st.session_state: st.session_state.reg_missing = []

    if not st.session_state.logged_in:
        _, center_col, _ = st.columns([1, 1.2, 1])
        with center_col:
            st.markdown("""
            <div style="text-align: center; margin-top: 20px; margin-bottom: 20px;">
                <span style="font-size: 42px;">🔐</span>
                <h2 style="margin: 8px 0 0 0; font-size: 24px; font-weight: 800;">ĐĂNG NHẬP HỆ THỐNG</h2>
                <p style="color: #888; font-size: 13px; margin-top: 4px;">Phần mềm Cụ thể hóa Văn bản Hành chính</p>
            </div>
            """, unsafe_allow_html=True)
            tab_login, tab_register, tab_forgot = st.tabs(["Đăng nhập", "Đăng ký tài khoản", "Quên mật khẩu"])
            with tab_login:
                with st.form("login_form"):
                    username = st.text_input("Tên đăng nhập")
                    password = st.text_input("Mật khẩu", type="password")
                    btn_login = st.form_submit_button("Đăng nhập", use_container_width=True)
                if btn_login:
                    if username == "admin" and password == "Adminai":
                        st.session_state.logged_in = True
                        st.session_state.user_info = {"username": "admin", "fullname": "Quản trị viên", "email_phone": "N/A"}
                        st.rerun()
                    else:
                        try:
                            df = pd.read_csv(SHEET_CSV_URL, skiprows=2)
                            df.columns = [c.strip() for c in df.columns]
                            user_match = df[(df['username'].astype(str) == username) & (df['password'].astype(str) == password)]
                            if not user_match.empty:
                                user_data = user_match.iloc[0]
                                st.session_state.logged_in = True
                                contact_val = str(user_data.get('email_phone', user_data.get('Email/SĐT', 'Chưa cập nhật')))
                                st.session_state.user_info = {"username": username, "fullname": user_data['fullname'], "email_phone": contact_val}
                                st.rerun()
                            else:
                                st.error("Tên đăng nhập hoặc mật khẩu không chính xác!")
                        except Exception:
                            st.error("Chưa thể kết nối đến dữ liệu tài khoản.")
        return False
    return True

st.set_page_config(page_title="Phần mềm Cụ thể hóa Văn bản Hành chính", page_icon="🏛️", layout="wide")
if not check_login(): st.stop()

SYMBOL_MAP = {"Kế hoạch": "KH", "Công văn": "CV", "Báo cáo": "BC", "Tờ trình": "TTr", "Thông báo": "TB", "Quyết định": "QĐ", "Hướng dẫn": "HD"}

TEMPLATES_CONFIG = {
    "Khối Đảng": {
        "Kế hoạch": "KẾ HOẠCH\nVề việc triển khai thực hiện...\n\nI. MỤC ĐÍCH, YÊU CẦU\n1. Mục đích\n2. Yêu cầu\n\nII. NỘI DUNG VÀ NHIỆM VỤ TRỌNG TÂM\n1. Công tác quán triệt, tuyên truyền\n2. Các nhiệm vụ và giải pháp cụ thể\n\nIII. TỔ CHỨC THỰC HIỆN\n1. Phân công trách nhiệm cho các cơ quan, đơn vị\n2. Chế độ thông tin, báo cáo",
        "Công văn": "Kính gửi: Các cơ quan, đơn vị trực thuộc.\n\n1. Căn cứ văn bản chỉ đạo của cấp trên về nhiệm vụ...\n2. Đề nghị các cơ quan, đơn vị tập trung triển khai thực hiện các nội dung trọng tâm sau đây...\n3. Giao bộ phận chuyên môn theo dõi, đôn đốc và tổng hợp báo cáo kết quả thực hiện định kỳ.",
        "Báo cáo": "BÁO CÁO\nKết quả thực hiện nhiệm vụ...\n\nI. TÌNH HÌNH VÀ KẾT QUẢ TRIỂN KHAI\n1. Công tác lãnh đạo, chỉ đạo và quán triệt\n2. Kết quả thực hiện các chỉ tiêu, nhiệm vụ trọng tâm\n\nII. ĐÁNH GIÁ CHUNG\n1. Ưu điểm, kết quả đạt được\n2. Tồn tại, hạn chế và nguyên nhân\n\nIII. NHIỆM VỤ VÀ GIẢI PHÁP TRỌNG TÂM THỜI GIAN TỚI",
        "Tờ trình": "TỜ TRÌNH\nVề việc...\n\nKính gửi: Ban Thường vụ Đảng ủy.\n\nI. SỰ CẦN THIẾT VÀ CĂN CỨ TRÌNH\nII. NỘI DUNG CHÍNH CỦA ĐỀ ÁN / KẾ HOẠCH\nIII. ĐỀ XUẤT, KIẾN NGHỊ",
        "Thông báo": "THÔNG BÁO\nKết luận của Thường trực Đảng ủy tại cuộc họp...\n\n1. Đánh giá khái quát tình hình triển khai công tác vừa qua\n2. Ý kiến chỉ đạo kết luận cụ thể đối với từng nhiệm vụ\n3. Giao trách nhiệm cho các ban đảng, cơ quan đơn vị triển khai thực hiện",
        "Quyết định": "QUYẾT ĐỊNH\nVề việc...\n\n- Căn cứ Điều lệ Đảng và Quy chế làm việc...\nBAN THƯỜNG VỤ QUYẾT ĐỊNH:\nĐiều 1. Phân công, giao nhiệm vụ cụ thể cho các tập thể, cá nhân.\nĐiều 2. Trách nhiệm thi hành của các cơ quan, đơn vị liên quan.\nĐiều 3. Quyết định có hiệu lực kể từ ngày ký.",
        "Hướng dẫn": "HƯỚNG DẪN\nVề việc thực hiện...\n\nI. MỤC ĐÍCH, YÊU CẦU\nII. ĐỐI TƯỢNG VÀ PHẠM VI ÁP DỤNG\nIII. NỘI DUNG VÀ QUY TRÌNH THỰC HIỆN\nIV. TỔ CHỨC THỰC HIỆN"
    },
    "Khối Nhà nước": {
        "Kế hoạch": "KẾ HOẠCH\nVề việc thực hiện...\n\nI. MỤC ĐÍCH, YÊU CẦU\nII. NHIỆM VỤ VÀ GIẢI PHÁP TRỌNG TÂM\nIII. KINH PHÍ VÀ TỔ CHỨC THỰC HIỆN",
        "Công văn": "Kính gửi: Các phòng, ban, đơn vị trực thuộc.\n\n1. Căn cứ các quy định pháp luật và chỉ đạo của cấp trên...\n2. Triển khai thực hiện các nhiệm vụ chuyên môn theo thẩm quyền...\n3. Yêu cầu các đơn vị nghiêm túc triển khai và báo cáo kết quả.",
        "Báo cáo": "BÁO CÁO\nTình hình thực hiện công tác...\n\nI. KẾT QUẢ THỰC HIỆN CÁC NHIỆM VỤ TRỌNG TÂM\nII. ĐÁNH GIÁ TỒN TẠI, HẠN CHẾ VÀ NGUYÊN NHÂN\nIII. PHƯƠNG HƯỚNG, NHIỆM VỤ TRONG THỜI GIAN TỚI",
        "Tờ trình": "TỜ TRÌNH\nVề việc phê duyệt...\n\nKính gửi: Ủy ban nhân dân.\n\nI. CĂN CỨ PHÁP LÝ VÀ SỰ CẦN THIẾT\nII. NỘI DUNG ĐỀ XUẤT\nIII. KIẾN NGHỊ VÀ PHƯƠNG ÁN GIẢI QUYẾT",
        "Thông báo": "THÔNG BÁO\nVề ý kiến kết luận của Lãnh đạo UBND tại cuộc họp...\n\n1. Nội dung đánh giá tình hình công tác\n2. Các ý kiến chỉ đạo, giao nhiệm vụ cụ thể cho các phòng ban\n3. Tiến độ hoàn thành và chế độ báo cáo",
        "Quyết định": "QUYẾT ĐỊNH\nVề việc...\n\n- Căn cứ Luật Tổ chức chính quyền địa phương...\nỦY BAN NHÂN DÂN QUYẾT ĐỊNH:\nĐiều 1. Quy định nội dung chi tiết.\nĐiều 2. Trách nhiệm tổ chức thực hiện.\nĐiều 3. Hiệu lực thi hành.",
        "Hướng dẫn": "HƯỚNG DẪN\nQuy trình thực hiện nhiệm vụ...\n\nI. MỤC ĐÍCH, YÊU CẦU\nII. ĐỐI TƯỢNG VÀ PHẠM VI ÁP DỤNG\nIII. NỘI DUNG VÀ QUY TRÌNH THỰC HIỆN\nIV. TỔ CHỨC THỰC HIỆN"
    }
}

def format_outline_to_html(text):
    lines = text.split('\n')
    formatted = ""
    for line in lines:
        l = line.strip()
        if not l: continue
        if re.match(r'^(I|II|III|IV|V|VI|VII|VIII)\.', l):
            formatted += f"<div style='font-weight: bold; margin-top: 8px; color: #63b3ed;'>{l}</div>"
        elif re.match(r'^\d+\.', l):
            formatted += f"<div style='margin-left: 15px; margin-top: 2px;'>{l}</div>"
        else:
            formatted += f"<div style='margin-left: 25px; font-style: italic; font-size: 0.95em;'>{l}</div>"
    return formatted

st.markdown("""
<style>
.app-header { background: linear-gradient(135deg, #7b0000 0%, #a81010 50%, #c41e1e 100%); border: 1px solid #e0a800; border-radius: 12px; padding: 20px 25px; margin-bottom: 20px; box-shadow: 0 4px 15px rgba(0, 0, 0, 0.4); display: flex; align-items: center; justify-content: space-between; }
.app-header-title { color: #ffffff; font-size: 24px; font-weight: bold; letter-spacing: 0.5px; text-shadow: 1px 1px 3px rgba(0,0,0,0.6); margin: 0; }
.app-header-sub { color: #ffd700; font-size: 13px; margin-top: 4px; font-weight: 500; }
.section-badge { background: linear-gradient(90deg, #d4af37 0%, #f3e5ab 100%); color: #4a2c00; font-size: 11px; font-weight: bold; padding: 3px 8px; border-radius: 4px; text-transform: uppercase; display: inline-block; margin-bottom: 6px; }
.a4-wrapper { background: radial-gradient(circle, #3d434d 0%, #20242b 100%); padding: 20px; border-radius: 10px; border: 1px solid #4a5568; display: flex; justify-content: center; width: 100%; box-shadow: inset 0 0 15px rgba(0,0,0,0.5); }
.a4-paper { background-color: #ffffff !important; color: #000000 !important; width: 100%; padding: 30px 35px; font-family: 'Times New Roman', Times, serif; font-size: 10.5pt; line-height: 1.35; box-shadow: 0 10px 25px rgba(0, 0, 0, 0.6); max-height: 620px; overflow-y: auto; box-sizing: border-box; border-radius: 2px; }
.header-table, .footer-table { width: 100%; border-collapse: collapse; margin-bottom: 10px; border: none !important; }
.header-table td, .footer-table td { vertical-align: top; font-family: 'Times New Roman', Times, serif; font-size: 10pt; line-height: 1.2; color: #000000; padding: 0px; }
.custom-underline { display: inline-block; border-bottom: 1px solid #000000; padding-bottom: 2px; line-height: 1.1; }
.title-block { text-align: center; font-weight: bold; font-size: 12pt; margin-top: 10px; margin-bottom: 4px; }
.trich-yeu-block { text-align: center; font-weight: bold; font-size: 11pt; margin-top: 4px; margin-bottom: 4px; }
.short-line { width: 35%; margin: 6px auto 12px auto; border: 0; border-top: 1px solid #000000; }
.content-para { text-align: justify; text-indent: 1cm; margin-bottom: 5px; line-height: 1.35; }
.heading-para { font-weight: bold; font-size: 10.5pt; margin-top: 10px; margin-bottom: 3px; }
.noi-nhan-block { padding-left: 1cm !important; text-align: left; font-size: 9.5pt; line-height: 1.2; }
.chat-user-box { background: linear-gradient(90deg, #242933 0%, #1a1e24 100%); color: #ffffff; padding: 12px 16px; border-radius: 8px; margin-bottom: 8px; display: flex; align-items: center; font-size: 13px; border-left: 3px solid #e53e3e; border: 1px solid #323946; }
.chat-user-icon { background: linear-gradient(135deg, #e53e3e 0%, #9b2c2c 100%); color: white; width: 30px; height: 30px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin-right: 12px; font-size: 14px; flex-shrink: 0; box-shadow: 0 2px 5px rgba(0,0,0,0.3); }
.chat-ai-box { background: linear-gradient(90deg, #1f2d24 0%, #17241c 100%); color: #ffffff; padding: 12px 16px; border-radius: 8px; margin-bottom: 16px; display: flex; align-items: center; font-size: 13px; border-left: 3px solid #38a169; border: 1px solid #234e32; }
.chat-ai-icon { background: linear-gradient(135deg, #38a169 0%, #22543d 100%); color: white; width: 30px; height: 30px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin-right: 12px; font-size: 14px; flex-shrink: 0; box-shadow: 0 2px 5px rgba(0,0,0,0.3); }
</style>
""", unsafe_allow_html=True)

with st.sidebar:
    st.markdown("""
    <div style="text-align: center; padding-bottom: 10px;">
        <span style="font-size: 32px;">🏛️</span>
        <h3 style="color: #ffd700; margin: 0; font-size: 18px;">HỆ THỐNG ĐIỀU HÀNH</h3>
        <p style="color: #a0aec0; font-size: 11px;">Chuẩn Thể Thức Đảng & Nhà Nước</p>
    </div>
    """, unsafe_allow_html=True)
    st.write("---")
    the_thuc = st.radio("Chọn Khối văn bản:", ["Khối Đảng", "Khối Nhà nước"])
    if the_thuc == "Khối Đảng":
        the_thuc_doc = "Hướng dẫn 05-HD/VPTW của Văn phòng Trung ương Đảng"
        default_agency = "ĐẢNG ỦY PHƯƠNG NHƠN TRẠCH"
    else:
        the_thuc_doc = "Nghị định 30/2020/NĐ-CP của Chính phủ"
        default_agency = "UBND PHƯƠNG NHƠN TRẠCH"

    st.info(f"📌 **Áp dụng:** {the_thuc_doc}")
    st.subheader("⚙️ Cấu hình AI")
    api_key = st.text_input("Gemini API key", value=config_data.get("gemini_key", ""), type="password")
    if st.button("💾 Lưu API Key vĩnh viễn", use_container_width=True):
        if save_config({"gemini_key": api_key}): st.success("Đã lưu!")
    
    st.write("---")
    user_info = st.session_state.get("user_info", {})
    with st.popover(f"👤 Tài khoản ({user_info.get('fullname', 'User')})"):
        if st.button("🚪 Đăng xuất", use_container_width=True):
            st.session_state.logged_in = False; st.rerun()

st.markdown('<div class="app-header"><div><div class="app-header-title">🏛️ PHẦN MỀM CỤ THỂ HÓA VĂN BẢN HÀNH CHÍNH</div><div class="app-header-sub">HỆ THỐNG HỖ TRỢ BIÊN SOẠN & XỬ LÝ VĂN KIỆN ĐẢNG - CHÍNH QUYỀN TỰ ĐỘNG BẰNG AI</div></div><div style="font-size: 38px;">🇻🇳</div></div>', unsafe_allow_html=True)

col1, col2 = st.columns([1, 1])
with col1:
    st.markdown('<span class="section-badge">BƯỚC 1</span> <b>File nguồn & Loại văn bản</b>', unsafe_allow_html=True)
    uploaded_files = st.file_uploader("Tải file nguồn/Đề cương (.docx, .pdf, .txt):", type=["pdf", "docx", "txt"], accept_multiple_files=True)
    loai_vb = st.selectbox("Chọn Loại văn bản đầu ra:", ["Kế hoạch", "Công văn", "Báo cáo", "Tờ trình", "Thông báo", "Quyết định", "Hướng dẫn"])

with col2:
    st.markdown('<span class="section-badge">BƯỚC 2</span> <b>Yêu cầu & Cơ quan ban hành</b>', unsafe_allow_html=True)
    yeu_cau = st.text_area("Anh muốn cụ thể hóa như thế nào?:", height=100)
    if "current_agency" not in st.session_state: st.session_state.current_agency = default_agency
    co_quan = st.text_input("Cơ quan ban hành dự thảo:", value=st.session_state.current_agency)

st.markdown('<br>', unsafe_allow_html=True)
col3_1, col3_2 = st.columns([1, 1])
with col3_1:
    st.markdown('<span class="section-badge">BƯỚC 3</span> <b>File mẫu riêng & Mẫu gợi ý chuẩn</b>', unsafe_allow_html=True)
    custom_template_file = st.file_uploader("Tải file mẫu riêng (.doc, .docx, .pdf):", type=["doc", "docx", "pdf"], key="custom_template")

with col3_2:
    st.markdown(f'<span style="color: #ffd700; font-size: 13px;">📌</span> <b>Mẫu gợi ý / Đề cương chuẩn ({the_thuc}):</b>', unsafe_allow_html=True)
    current_default_outline = TEMPLATES_CONFIG[the_thuc].get(loai_vb, "")
    selected_builtin = st.selectbox("Mẫu gợi ý / Đề cương chuẩn:", [f"📌 Đề cương {loai_vb} chuẩn", "(Không chọn mẫu gợi ý)"], label_visibility="collapsed")
    if selected_builtin != "(Không chọn mẫu gợi ý)":
        st.markdown(f'<div style="background-color: #161c24; border: 1px solid #2b6cb0; border-radius: 6px; padding: 15px;">{format_outline_to_html(current_default_outline)}</div>', unsafe_allow_html=True)

st.markdown("---")
btn_process = st.button("⚡ PHÂN TÍCH & CỤ THỂ HÓA VĂN BẢN", type="primary", use_container_width=True)

if "draft_text" not in st.session_state: st.session_state.draft_text = ""
if "chat_history" not in st.session_state: st.session_state.chat_history = []

def call_gemini(prompt_text, parts=None):
    genai.configure(api_key=api_key)
    # Dùng chuẩn định danh chung hoạt động ổn định nhất qua API
    model = genai.GenerativeModel("gemini-2.5-flash")
    if parts:
        response = model.generate_content([prompt_text] + parts)
    else:
        response = model.generate_content(prompt_text)
    return response.text

if btn_process:
    if not api_key:
        st.error("Vui lòng nhập Gemini API Key ở cột bên trái!")
    elif not uploaded_files and not yeu_cau:
        st.warning("Vui lòng tải file nguồn hoặc nhập yêu cầu!")
    else:
        with st.spinner("Đang phân tích dữ liệu và tổng hợp văn bản chuẩn thể thức..."):
            try:
                content_parts, extracted_texts = [], []
                
                for uf in uploaded_files:
                    bytes_data = uf.read()
                    if uf.name.lower().endswith('.pdf'):
                        try:
                            reader = pypdf.PdfReader(io.BytesIO(bytes_data))
                            pdf_text = "".join([page.extract_text() or "" for page in reader.pages])
                            if len(pdf_text.strip()) > 50: extracted_texts.append(f"--- NỘI DUNG FILE NGUỒN {uf.name} ---\n" + pdf_text)
                            else: content_parts.append({"mime_type": "application/pdf", "data": bytes_data})
                        except Exception: content_parts.append({"mime_type": "application/pdf", "data": bytes_data})
                    elif uf.name.lower().endswith('.docx'):
                        doc_file = docx.Document(io.BytesIO(bytes_data))
                        extracted_texts.append(f"--- NỘI DUNG FILE NGUỒN {uf.name} ---\n" + "\n".join([p.text for p in doc_file.paragraphs]))
                    else: content_parts.append({"mime_type": uf.type, "data": bytes_data})
                
                custom_template_parts = []
                if custom_template_file is not None:
                    tpl_bytes = custom_template_file.read()
                    mime_type = custom_template_file.type if custom_template_file.type else "application/msword"
                    if custom_template_file.name.lower().endswith('.pdf'): mime_type = "application/pdf"
                    elif custom_template_file.name.lower().endswith('.docx'): mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    custom_template_parts.append({"mime_type": mime_type, "data": tpl_bytes})

                outline_prompt = ""
                if custom_template_file is not None:
                    outline_prompt = """
                    BẮT BUỘC TUÂN THỦ TUYỆT ĐỐI CẤU TRÚC VÀ KHUNG SƯỞN CỦA FILE MẪU VĂN BẢN RIÊNG ĐƯỢC ĐÍNH KÈM:
                    - Hãy phân tích bố cục, hệ thống các mục (I, II, III...), cách phân chia đề mục của file mẫu riêng để làm khuôn mẫu cấu trúc.
                    - LẤY NỘI DUNG TỪ TÀI LIỆU CẤP TRÊN ĐỂ VIẾT VÀO CÁC MỤC ĐÓ, TUYỆT ĐỐI KHÔNG COPY NỘI DUNG CŨ CỦA FILE MẪU VÀO DỰ THẢO.
                    """
                elif selected_builtin != "(Không chọn mẫu gợi ý)":
                    outline_prompt = f"BẮT BUỘC ÁP DỤNG ĐỀ CƯƠNG CHUẨN (HD 05 / NĐ 30):\n{current_default_outline}"
                else:
                    outline_prompt = "Áp dụng cấu trúc chuẩn hành chính."

                rule_doc_type = "ĐÂY LÀ CÔNG VĂN:\n1. Không viết tiêu đề CÔNG VĂN giữa trang.\n2. Trích yếu: V/v...\n3. Kính gửi:..." if loai_vb == "Công văn" else f"ĐÂY LÀ {loai_vb.upper()}:\n1. Tên loại viết hoa giữa trang.\n2. Trích yếu nội dung."

                prompt = f"""
                Bạn là chuyên gia soạn thảo văn bản hành chính Việt Nam. Hãy soạn thảo 01 dự thảo văn bản hoàn chỉnh, chuẩn xác theo thể thức nhà nước.
                THỂ THỨC: {the_thuc} | CƠ QUAN: {co_quan} | LOẠI: {loai_vb}
                YÊU CẦU: {yeu_cau}
                
                {outline_prompt}
                
                DỮ LIỆU TÀI LIỆU NGUỒN CỦA CẤP TRÊN:
                {"".join(extracted_texts)}
                
                QUY TẮC BẮT BUỘC:
                {rule_doc_type}
                - Không viết Quốc hiệu, Tiêu ngữ, Tên cơ quan, Số/Ký hiệu, Ngày tháng ở đầu bài.
                - Không dùng ký tự Markdown (*, #, _).
                - Ở mục cuối cùng bắt buộc ghi đúng định dạng:
                  Nơi nhận:
                  - Như trên;
                  - Lưu: VP.
                  T/M BAN THƯỜNG VỤ
                  BÍ THƯ
                  (Để trống khoảng ký tên)
                  Họ và Tên
                """
                
                final_parts = content_parts + custom_template_parts
                response_text = call_gemini(prompt, final_parts if final_parts else None)
                st.session_state.draft_text = response_text
                st.session_state.current_agency = co_quan
                st.session_state.chat_history = []
                st.success("Đã cụ thể hóa văn bản thành công!")
            except Exception as e:
                st.error(f"Lỗi xử lý Gemini API: {str(e)}")

# HIỂN THỊ TRANG A4 VÀ CHAT SỬA ĐỔI
if st.session_state.draft_text:
    res_col1, res_col2 = st.columns([1.2, 0.8])
    with res_col1:
        st.markdown('##### 📄 BẢN DỰ THẢO VĂN BẢN (A4)')
        clean_text = re.sub(r'[\*#_]', '', st.session_state.draft_text)
        raw_lines = [l.strip() for l in clean_text.split('\n') if l.strip()]
        
        filtered_lines = []
        for l in raw_lines:
            l_up = l.upper()
            if any(k in l_up for k in ["ĐẢNG BỘ", "ĐẢNG CỘNG SẢN", "CỘNG HÒA XÃ HỘI", "ĐỘC LẬP - TỰ DO"]) and len(l) < 80: continue
            if re.search(r',\s*NGÀY\s+.*\s+THÁNG\s+.*\s+NĂM', l_up) and len(l) < 80: continue
            if ("UBND" in l_up or "ĐẢNG ỦY" in l_up) and len(l) < 60 and not l_up.startswith("KẾ HOẠCH") and not l_up.startswith("CÔNG VĂN"): continue
            if re.match(r'^SỐ\s*:', l_up) or re.match(r'^SỐ\s*-', l_up): continue
            filtered_lines.append(l)

        body_lines, noi_nhan_list = [], ["- Như trên;", "- Lưu: VP."]
        chuc_vu_signer = "T/M BAN THƯỜNG VỤ\nBÍ THƯ" if the_thuc == "Khối Đảng" else "TM. ỦY BAN NHÂN DÂN\nCHỦ TỊCH"
        ten_signer, in_footer = "Họ và Tên", False
        
        for l in filtered_lines:
            l_strip = l.strip()
            l_up = l_strip.upper()
            if l_up.startswith("NƠI NHẬN:") or l_up == "NƠI NHẬN": in_footer = True; continue
            if in_footer:
                if l_strip.startswith("-"):
                    if l_strip not in noi_nhan_list: noi_nhan_list.append(l_strip)
                elif any(k in l_up for k in ["T/M", "TM.", "BÍ THƯ", "CHỦ TỊCH", "PHÓ BÍ THƯ"]): chuc_vu_signer = l_strip
                elif len(l_strip) < 40 and not l_strip.startswith("I") and not l_strip.startswith("1") and l_strip.lower() != "họ và tên": ten_signer = l_strip
            else:
                if any(k in l_up for k in ["T/M ", "TM. ", "BÍ THƯ", "CHỦ TỊCH"]) and len(l_strip) < 50: in_footer = True; chuc_vu_signer = l_strip
                else: body_lines.append(l_strip)

        trich_yeu_cv = ""
        if loai_vb == "Công văn" and body_lines:
            if body_lines[0].startswith("V/v") or body_lines[0].startswith("Về việc"): trich_yeu_cv = body_lines.pop(0)

        agency_display = st.session_state.get("current_agency", co_quan)
        dia_danh = "Nhơn Trạch"
        if "ĐẠI PHƯỚC" in agency_display.upper(): dia_danh = "Đại Phước"
        elif "NHƠN TRẠCH" in agency_display.upper(): dia_danh = "Nhơn Trạch"

        type_code = SYMBOL_MAP.get(loai_vb, "CV")
        so_ky_hieu = f"-{type_code}/ĐU" if the_thuc == "Khối Đảng" else f"/{type_code}-UBND"

        sub_cv = f"<br><br><i>{trich_yeu_cv}</i>" if trich_yeu_cv else ""
        if the_thuc == "Khối Đảng":
            header_table = f'<table class="header-table"><tr><td style="width: 48%; text-align: center;"><b>ĐẢNG BỘ THÀNH PHỐ ĐỒNG NAI</b><br><b>{agency_display.upper()}</b><br><span style="font-size: 7pt;">*</span><br>Số: &nbsp;&nbsp;&nbsp;&nbsp;{so_ky_hieu}{sub_cv}</td><td style="width: 52%; text-align: center;"><span class="custom-underline"><b>ĐẢNG CỘNG SẢN VIỆT NAM</b></span><br><br><i>{dia_danh}, ngày &nbsp;&nbsp;&nbsp; tháng 8 năm 2026</i></td></tr></table>'
        else:
            header_table = f'<table class="header-table"><tr><td style="width: 45%; text-align: center;">UBND THÀNH PHỐ ĐỒNG NAI<br><b><span class="custom-underline">{agency_display.upper()}</span></b><br><span style="font-size: 7pt;">*</span><br>Số: &nbsp;&nbsp;&nbsp;&nbsp;{so_ky_hieu}{sub_cv}</td><td style="width: 55%; text-align: center;"><b>CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM</b><br><b><span class="custom-underline">Độc lập - Tự do - Hạnh phúc</span></b><br><i>{dia_danh}, ngày &nbsp;&nbsp;&nbsp; tháng 8 năm 2026</i></td></tr></table>'

        body_content, is_trich_yeu = "", False
        for line in body_lines:
            if re.match(r'^(I|II|III|IV|V|VI|VII|VIII)\.', line): body_content += f'<div class="heading-para">{line}</div>'; is_trich_yeu = False
            elif re.match(r'^\d+\.', line) and len(line) < 80: body_content += f'<div class="heading-para">{line}</div>'; is_trich_yeu = False
            elif line.isupper() and len(line) < 100 and loai_vb != "Công văn": body_content += f'<div class="title-block">{line}</div>'; is_trich_yeu = True
            elif is_trich_yeu and loai_vb != "Công văn": body_content += f'<div class="trich-yeu-block">{line}</div><hr class="short-line">'; is_trich_yeu = False
            elif line.startswith("Kính gửi:") or line.startswith("-"): body_content += f'<div style="text-align: left; margin-bottom: 4px; padding-left: 10px;">{line}</div>'; is_trich_yeu = False
            else: body_content += f'<div class="content-para">{line}</div>'; is_trich_yeu = False

        footer_table = f"""
        <table class="footer-table" style="margin-top: 20px;">
            <tr>
                <td style="width: 45%; text-align: left; vertical-align: top;" class="noi-nhan-block">
                    <b><u>Nơi nhận:</u></b><br>{"<br>".join(noi_nhan_list)}
                </td>
                <td style="width: 55%; text-align: center; vertical-align: top;">
                    <b>{chuc_vu_signer.replace(chr(10), "<br>")}</b><br><br><br><br><br><b>{ten_signer}</b>
                </td>
            </tr>
        </table>
        """
        st.markdown(f'<div class="a4-wrapper"><div class="a4-paper">{header_table}{body_content}{footer_table}</div></div>', unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

        def generate_docx(b_lines, agency_name, form_type, doc_type_str, cv_subj, n_nhan, c_vu, t_ky, dia_danh_str):
            doc = docx.Document()
            for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(2), Cm(2), Cm(3), Cm(2)
            table = doc.add_table(rows=1, cols=2)
            table.alignment, table.autofit = WD_TABLE_ALIGNMENT.CENTER, False
            c_l, c_r = table.cell(0, 0), table.cell(0, 1)
            c_l.width, c_r.width = Cm(8.5), Cm(8.5)
            
            t_code = SYMBOL_MAP.get(doc_type_str, "CV")
            code_str = f"-{t_code}/ĐU" if form_type == "Khối Đảng" else f"/{t_code}-UBND"
            
            p_l = c_l.paragraphs[0]; p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if form_type == "Khối Đảng":
                p_l.add_run("ĐẢNG BỘ THÀNH PHỐ ĐỒNG NAI\n").font.size = Pt(12)
                r2 = p_l.add_run(f"{agency_name.upper()}\n"); r2.font.bold = True; r2.font.size = Pt(12)
                p_l.add_run("*\n").font.size = Pt(9)
                p_l.add_run(f"Số:       {code_str}").font.size = Pt(12)
                if cv_subj: p_l.add_run(f"\n\n{cv_subj}").font.italic = True
            else:
                p_l.add_run("UBND THÀNH PHỐ ĐỒNG NAI\n").font.size = Pt(12)
                r2 = p_l.add_run(f"{agency_name.upper()}\n"); r2.font.bold = True; r2.font.underline = True; r2.font.size = Pt(12)
                p_l.add_run("*\n").font.size = Pt(9)
                p_l.add_run(f"Số:       {code_str}").font.size = Pt(12)
                if cv_subj: p_l.add_run(f"\n\n{cv_subj}").font.italic = True

            p_r = c_r.paragraphs[0]; p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if form_type == "Khối Đảng":
                r1 = p_r.add_run("ĐẢNG CỘNG SẢN VIỆT NAM"); r1.font.bold = True; r1.font.underline = True; r1.font.size = Pt(12)
                p_r.add_run(f"\n\n{dia_danh_str}, ngày     tháng 8 năm 2026").font.italic = True
            else:
                p_r.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n").font.bold = True
                r2 = p_r.add_run("Độc lập - Tự do - Hạnh phúc"); r2.font.bold = True; r2.font.underline = True; r2.font.size = Pt(12.5)
                p_r.add_run(f"\n{dia_danh_str}, ngày     tháng 8 năm 2026").font.italic = True

            doc.add_paragraph()
            for line in b_lines:
                p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(4)
                if re.match(r'^(I|II|III|IV|V|VI|VII|VIII)\.', line): p.add_run(line).font.bold = True
                elif re.match(r'^\d+\.', line) and len(line) < 80: p.add_run(line).font.bold = True
                elif line.startswith("Kính gửi:"): p.add_run(line)
                else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent = Cm(1.27); p.add_run(line)

            t_foot = doc.add_table(rows=1, cols=2); t_foot.alignment, t_foot.autofit = WD_TABLE_ALIGNMENT.CENTER, False
            c_fl, c_fr = t_foot.cell(0, 0), t_foot.cell(0, 1)
            c_fl.width, c_fr.width = Cm(8.5), Cm(8.5)
            p_fl = c_fl.paragraphs[0]; p_fl.paragraph_format.left_indent = Cm(1)
            r_nn = p_fl.add_run("Nơi nhận:\n"); r_nn.font.bold = True; r_nn.font.underline = True; r_nn.font.italic = True
            for item in n_nhan: p_fl.add_run(f"{item}\n")
            
            p_fr = c_fr.paragraphs[0]; p_fr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cv = p_fr.add_run(f"{c_vu}\n\n\n\n\n"); r_cv.font.bold = True
            p_fr.add_run(t_ky).font.bold = True
            
            bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

        st.download_button(
            label="📥 TẢI VỀ FILE WORD (.DOCX)",
            data=generate_docx(body_lines, agency_display, the_thuc, loai_vb, trich_yeu_cv, noi_nhan_list, chuc_vu_signer, ten_signer, dia_danh),
            file_name=f"Du_Thao_{loai_vb}_{dia_danh}.docx".replace(" ", "_"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )

    with res_col2:
        st.markdown('##### 💬 TRỢ LÝ AI CHỈNH SỬA')
        edit_instruction = st.text_area("Nhập yêu cầu chỉnh sửa...", height=120, label_visibility="collapsed", placeholder="Nhập yêu cầu chỉnh sửa văn bản...")
        
        if st.button("Chỉnh sửa dự thảo", use_container_width=True):
            if not api_key:
                st.error("Vui lòng nhập Gemini API Key ở cột bên trái!")
            elif edit_instruction:
                with st.spinner("Đang cập nhật lại dự thảo..."):
                    try:
                        res_edit_text = call_gemini(f"BẢN DỰ THẢO:\n{st.session_state.draft_text}\n\nYÊU CẦU CHỈNH SỬA: {edit_instruction}\nHãy cập nhật lại toàn bộ văn bản theo yêu cầu. Không dùng ký tự Markdown.")
                        full_res = res_edit_text
                        
                        edit_lower = edit_instruction.lower()
                        if "đại phước" in edit_lower:
                            st.session_state.current_agency = st.session_state.current_agency.replace("NHƠN TRẠCH", "ĐẠI PHƯỚC").replace("Nhơn Trạch", "Đại Phước")
                        elif "nhơn trạch" in edit_lower:
                            st.session_state.current_agency = st.session_state.current_agency.replace("ĐẠI PHƯỚC", "NHƠN TRẠCH").replace("Đại Phước", "Nhơn Trạch")
                            
                        st.session_state.draft_text = full_res
                        st.session_state.chat_history.append(edit_instruction)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Lỗi chỉnh sửa: {str(e)}")

        if st.session_state.chat_history:
            st.markdown("<br>", unsafe_allow_html=True)
            for cmd in reversed(st.session_state.chat_history):
                st.markdown(f'<div class="chat-user-box"><div class="chat-user-icon">⏰</div><div>{cmd}</div></div>', unsafe_allow_html=True)
                st.markdown('<div class="chat-ai-box"><div class="chat-ai-icon">📊</div><div><b>✅ Đã cập nhật văn bản lên trang Word!</b></div></div>', unsafe_allow_html=True)
